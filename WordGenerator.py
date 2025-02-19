import os
import re
import tempfile
from openai import OpenAI
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from PIL import Image
from io import BytesIO


# 1. Configure l'API OpenAI avec ta clé API
client = OpenAI()

# Demander à l'utilisateur de spécifier les noms des dossiers
input_dir_name = input("Indiquez le nom du dossier contenant les fichiers à traiter (par défaut : Lists) : ") or "Lists"
output_dir_name = input("Indiquez le nom du dossier où les fichiers modifiés seront sauvegardés (par défaut : OutputGPT) : ") or "OutputGPT"

# Construire les chemins complets en ajoutant "./"
input_dir = f"./{input_dir_name}"
output_dir = f"./{output_dir_name}"

# S'assurer que le dossier de sortie existe
os.makedirs(output_dir, exist_ok=True)

# Fonction pour générer les questions
def generate_questions(word):
    interrogative = f"Is there a {word} in the picture?"
    affirmative = f"There is a {word} in the picture."
    negative = f"There is no {word} in the picture."
    return interrogative, affirmative, negative

# Appliquer la même police à tous les runs dans un paragraphe
def apply_font_to_paragraph(paragraph, font_name, font_size):
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.bold = True
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), font_name)

# Fonction pour extraire le numéro de la liste depuis le nom du fichier
def extract_list_number(filename):
    match = re.search(r'List (\d+)', filename)
    return int(match.group(1)) if match else float('inf')

# Fonction pour extraire une image d'une cellule, si disponible
def extract_image_from_cell(cell, doc):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            drawing_elements = run.element.xpath('.//w:drawing')
            if drawing_elements:
                blip_elements = drawing_elements[0].xpath('.//a:blip')
                if blip_elements:  # Vérification si l'élément blip existe
                    blip = blip_elements[0]
                    rId = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                    return doc.part.related_parts[rId]
    return None


# Fonction pour générer les questions avec ChatGPT en utilisant la nouvelle API
def generate_questions_with_chatgpt(word):
    messages = [
        {"role": "system", "content": "You are a helpful assistant that creates educational content for 9-year-old children."},
        {
            "role": "user",
            "content": (
                f"Generate one interrogative, one affirmative, and one negative sentence using the word '{word}'. "
                "The sentences should be in simple past tense, each containing 8 to 10 words, and suitable for a 10-year-old child. "
                "The sentences should vary using 'to be', 'to have', 'to go', and interrogative words like what, where, how many, why, who, when, how, etc."
            ),
        },
    ]

    response = client.chat.completions.create(
        model="gpt-3.5-turbo",  # Remplace par "gpt-4o-mini" si tu utilises GPT-4
        messages=messages,
        max_tokens=100,
        n=1,
        temperature=0.7,
    )

    # Récupérer la réponse du modèle
    message = response.choices[0].message.content.strip()
  
    # Diviser la réponse en lignes et enlever les titres
    sentences = message.split("\n")

    cleaned_sentences = []
    
    for sentence in sentences:
        # Enlever les titres et lignes vides
        if sentence.lower().startswith(('interrogative', 'affirmative', 'negative')):
            # Si la phrase est dans le format 'Interrogative: What is ...'
            cleaned_sentence = sentence.split(":", 1)[1].strip()
            cleaned_sentences.append(cleaned_sentence)
        else:
            cleaned_sentence = sentence.strip("- ").strip()
            if cleaned_sentence:
                cleaned_sentences.append(cleaned_sentence)
                
    cleaned_sentences = [sentence for sentence in cleaned_sentences if sentence]

    # S'assurer que nous avons bien trois phrases
    return cleaned_sentences[0], cleaned_sentences[1], cleaned_sentences[2]
# Fonction pour extraire le mot directement sous une image dans une cellule
def extract_word_from_cell(cell):
    lines = [line.strip() for line in cell.text.splitlines() if line.strip()]
    if lines:
        first_line = lines[0].strip()
        # Si la première ligne est de la forme "1. Kitchen" ou "1. To Kick"
        match = re.match(r'\d+\.\s*(.*)', first_line)
        if match:
            word = match.group(1).strip()
            return word  # Renvoie "To Kick" ou "Kitchen" sans le "1."
        
        # Si le mot commence par "To " sans numéro devant
        if first_line.lower().startswith("to "):
            return first_line  # Prend toute la ligne si c'est "To + verb"
        
    return None  # Aucun mot valide trouvé

# Fonction pour extraire le nom du fichier jusqu'au tiret '-'
def extract_name_before_dash(filename):
    # Séparer les parties du nom de fichier par le tiret
    parts = filename.split('-')[0].strip()
    # Séparer les mots par des espaces et capitaliser chaque mot
    name_parts = [word.title() for word in parts.split()]
    # Rejoindre les mots en un seul nom
    nameFile = " ".join(name_parts)
    return nameFile


# Fonction modifiée pour ajouter une image avec vérification des DPI
# Fonction modifiée pour ajouter une image avec vérification des DPI
def add_image_with_fallback(run, image_part, width=Inches(0.75)):
    with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp:
        tmp.write(image_part.blob)
        tmp.seek(0)
        
        # Utilisation de Pillow pour ouvrir l'image et vérifier les DPI
        with Image.open(tmp.name) as img:
            dpi = img.info.get('dpi', (96, 96))  # Récupérer le DPI ou définir une valeur par défaut
            horz_dpi, vert_dpi = dpi
            if horz_dpi == 0 or vert_dpi == 0:
                horz_dpi = 96  # Valeur par défaut pour horz_dpi
                vert_dpi = 96  # Valeur par défaut pour vert_dpi
            
            # Conversion en pixels
            width_px = int((width / Inches(1)) * horz_dpi)
            
            # Redimensionner l'image en utilisant LANCZOS
            img = img.resize((width_px, int(img.height * (width_px / img.width))), Image.LANCZOS)
            img.save(tmp.name)
        
        run.add_picture(tmp.name, width=width)

# Parcourir tous les fichiers dans le dossier Lists, triés par numéro de liste
for filename in sorted(os.listdir(input_dir), key=extract_list_number):
    if filename.endswith(".docx"):
        
        # Début du traitement de chaque fichier
        print("\n" + "="*50)
        print(f"Traitement du fichier : {filename}")
        print("="*50)

        # Récupérer le nom jusqu'au premier tiret '-'
        nameFile = extract_name_before_dash(filename)
        
        # Récupérer le numéro de la liste à partir du nom du fichier
        file_Index = extract_list_number(filename)
        if file_Index is None:
            print(f"[WARNING] Le fichier '{filename}' n'a pas de numéro de liste identifiable. Il est ignoré.")
            continue
        
        # Confirmation de la découverte du document
        print(f"[INFO] Document '{filename}' trouvé !")
        
        # Ouverture du fichier
        input_path = os.path.join(input_dir, filename)
        doc1 = Document(input_path)
        print(f"[INFO] Ouverture du fichier '{filename}' en cours...")

        # Accéder aux tableaux (deuxième tableau pour les mots et les images, premier tableau pour les images de secours)
        image_table = doc1.tables[0]
        text_table = doc1.tables[1]

        # Extraction des images et mots
        print("[INFO] Extraction des images et mots en cours...")
        words_with_images = []

        for i, row in enumerate(text_table.rows):
            for j, cell in enumerate(row.cells):
                word = extract_word_from_cell(cell)
                image_part = extract_image_from_cell(cell, doc1)                    
                if not image_part and i < len(image_table.rows):  # Si aucune image n'est trouvée dans le deuxième tableau
                    image_part = extract_image_from_cell(image_table.rows[i].cells[j], doc1)

                if word:
                    if not image_part:  # Si aucune image n'est trouvée dans les deux tableaux
                        print(f"[WARNING] Aucune image trouvée pour '{word}' dans le fichier '{filename}', à la ligne {i+1}. La cellule sera laissée vide.")
                    
                    words_with_images.append((word, image_part))

        # Confirmation de l'extraction terminée
        print("[INFO] Extraction des images et des mots terminée avec succès !")

        # Ouverture du template
        output_doc = Document("Basic vocabulary -  TEMPLATE.docx")
        print("[INFO] Ouverture du template en cours...") 
        
        # Modifier la phrase en haut du document et appliquer le formatage
        paragraph = output_doc.paragraphs[0]
        for paragraph in output_doc.paragraphs:
            if 'Basic vocabulary - List 3' in paragraph.text:
                run = paragraph.clear()  # Efface le texte existant pour le remplacer
                run = paragraph.add_run(f"{nameFile} (Simple past) - List {file_Index}")
                apply_font_to_paragraph(paragraph, "Goudy Old Style", 16)
                run.font.color.rgb = RGBColor(112, 48, 160)  # Appliquer la couleur #7030A0

        # Accéder aux tableaux du modèle
        output_table_1 = output_doc.tables[0]
        output_table_2 = output_doc.tables[1]
        print("[INFO] Insertion des images et génération des phrases en cours...") 
        
        # Remplir le premier tableau avec les mots de vocabulaire et les images
        for i, (word, image_part) in enumerate(words_with_images):
            row_index = i + 1  # Les lignes commencent à 1 (la première ligne est l'en-tête)
            row = output_table_1.rows[row_index]
            row2 = output_table_2.rows[row_index]

            # Générer les questions
            interrogative, affirmative, negative = generate_questions_with_chatgpt(word)
    
            # Insérer les phrases dans les colonnes correspondantes
            row.cells[2].text = interrogative  # Question (?)
            row.cells[3].text = affirmative    # Affirmative (Yes)
            row.cells[4].text = negative       # Negative (No)

            # Appliquer le formatage à chaque cellule du tableau
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    apply_font_to_paragraph(paragraph, "Goudy Old Style", 11)

            # Insérer l'image dans la première colonne si elle existe
            if image_part:
                paragraph = row.cells[1].paragraphs[0]
                run = paragraph.add_run()
                add_image_with_fallback(run, image_part)
                paragraph = row2.cells[1].paragraphs[0]
                run = paragraph.add_run()
                add_image_with_fallback(run, image_part)

        # Sauvegarde du fichier complété
        output_path = os.path.join(output_dir, f"{filename}")
        output_doc.save(output_path)
        print(f"[SUCCESS] Le fichier '{filename}' a été rempli avec succès et sauvegardé sous '{output_dir_name}'.")
        print("="*50 + "\n")
